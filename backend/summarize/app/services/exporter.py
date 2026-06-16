import io
import re
from datetime import datetime, date
from typing import Any

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def format_date(date_val: Any) -> str:
    """Format datetime/date objects or ISO strings into readable local date formats."""
    if not date_val:
        return "Not specified"
    try:
        if isinstance(date_val, (datetime, date)):
            d = date_val
        else:
            # Parse from ISO string, converting trailing Z for timezone compliance
            cleaned = str(date_val).replace("Z", "+00:00")
            try:
                d = datetime.fromisoformat(cleaned)
            except ValueError:
                # Handle dates like '2026-06-08'
                d = date.fromisoformat(cleaned)
                
        if isinstance(d, datetime):
            return d.strftime("%B %d, %Y, %I:%M %p")
        return d.strftime("%B %d, %Y")
    except Exception:
        return str(date_val)


def get_meeting_type_label(m_type: str | None) -> str:
    """Map the database meeting type string to user-friendly titles."""
    labels = {
        "team": "Team Meeting",
        "one-on-one": "1-on-1",
        "client": "Client Meeting",
        "standup": "Standup",
        "project-review": "Project Review",
        "brainstorm": "Brainstorming",
        "interview": "Interview",
        "training": "Training",
        "other": "Other",
    }
    return labels.get(str(m_type).lower(), str(m_type) if m_type else "Not specified")


def sanitize_text(text: str | None) -> str:
    """Clean text by stripping non-printable control parameters."""
    if not text:
        return ""
    # Strip ASCII control sequences
    sanitized = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", str(text))
    return sanitized.replace("\uFFFD", "").strip()


def clean_markdown(text: str | None) -> str:
    """Remove markdown symbols (**bold**, *italic*, ### headers, etc.) to produce clean text."""
    if not text:
        return ""
    # Remove bold markdown: **text** -> text
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    # Remove headers: ### Header -> Header
    text = re.sub(r"^\s*#+\s+", "", text, flags=re.MULTILINE)
    # Convert/clean bullet list symbols: * bullets, + bullets, - bullets -> • bullets
    text = re.sub(r"^\s*[-*+]\s+", "• ", text, flags=re.MULTILINE)
    # Remove italic markdown
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text.strip()


def validate_summary_data(summary_data: dict[str, Any]) -> dict[str, Any]:
    """Validate summary dictionaries ensuring required keys exist."""
    if not summary_data or not isinstance(summary_data, dict):
        raise ValueError("Invalid summary data: data is required")
        
    summary_text = summary_data.get("summary")
    if not summary_text or not str(summary_text).strip():
        raise ValueError("Invalid summary data: summary text is required")
        
    return {
        "meetingTitle": summary_data.get("meetingTitle") or "Untitled Meeting",
        "meetingDate": summary_data.get("meetingDate") or None,
        "meetingType": summary_data.get("meetingType") or None,
        "participants": list(summary_data.get("participants") or []),
        "location": summary_data.get("location") or None,
        "tags": list(summary_data.get("tags") or []),
        "summary": str(summary_text).strip(),
        "actionItems": list(summary_data.get("actionItems") or []),
        "decisions": list(summary_data.get("decisions") or []),
        "deadlines": list(summary_data.get("deadlines") or []),
        "extractedParticipants": list(summary_data.get("extractedParticipants") or []),
    }


def create_meeting_notes_template(summary_data: dict[str, Any]) -> dict[str, Any]:
    """Compile raw summary fields merging invitee list and generated participant names."""
    participants = list(summary_data.get("participants") or [])
    extracted = list(summary_data.get("extractedParticipants") or [])
    
    # Merge unique names
    all_participants = list(participants)
    for p in extracted:
        if p not in all_participants:
            all_participants.append(p)
            
    return {
        "title": clean_markdown(sanitize_text(summary_data.get("meetingTitle") or "Untitled Meeting")),
        "date": format_date(summary_data.get("meetingDate")),
        "type": get_meeting_type_label(summary_data.get("meetingType")),
        "participants": [clean_markdown(sanitize_text(p)) for p in all_participants],
        "location": clean_markdown(sanitize_text(summary_data.get("location") or "Not specified")),
        "tags": [clean_markdown(sanitize_text(t)) for t in (summary_data.get("tags") or [])],
        "summary": clean_markdown(sanitize_text(summary_data.get("summary"))),
        "actionItems": summary_data.get("actionItems") or [],
        "decisions": summary_data.get("decisions") or [],
        "deadlines": summary_data.get("deadlines") or [],
    }


def draw_first_page(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 9)
    canvas.setFillColor(colors.HexColor("#64748b")) # Slate 500
    canvas.drawString(50, 30, "Generated by MeetNote AI")
    canvas.drawRightString(545.27, 30, f"Page {canvas.getPageNumber()}")
    canvas.restoreState()


def draw_later_pages(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 9)
    canvas.setFillColor(colors.HexColor("#64748b")) # Slate 500
    canvas.drawString(50, 30, "Generated by MeetNote AI")
    canvas.drawRightString(545.27, 30, f"Page {canvas.getPageNumber()}")
    # Running header
    canvas.setStrokeColor(colors.HexColor("#e2e8f0"))
    canvas.setLineWidth(0.5)
    canvas.line(50, 800, 545.27, 800)
    canvas.drawString(50, 805, doc.title or "Meeting Notes")
    canvas.drawRightString(545.27, 805, "MeetNote")
    canvas.restoreState()


def generate_pdf(summary_data: dict[str, Any]) -> bytes:
    """Generate a formatted PDF document using ReportLab SimpleDocTemplate flowables."""
    validated = validate_summary_data(summary_data)
    template = create_meeting_notes_template(validated)
    
    buffer = io.BytesIO()
    
    # A4 dimensions (595.27 x 841.89 points) with 50-point margins
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=50,
        rightMargin=50,
        topMargin=50,
        bottomMargin=50,
    )
    
    doc.title = template["title"]
    doc.author = "MeetNote Meeting Notes"
    doc.creator = "MeetNote Summarizer"
    
    title_style = ParagraphStyle(
        name="DocTitle",
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0f172a"), # Slate 900
        spaceAfter=15,
        alignment=0, # Left-aligned
    )
    
    h1_style = ParagraphStyle(
        name="DocHeading1",
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#059669"), # Emerald 600
        spaceBefore=16,
        spaceAfter=8,
        keepWithNext=True,
    )
    
    body_style = ParagraphStyle(
        name="DocBody",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#334155"), # Slate 700
        spaceAfter=8,
    )
    
    bullet_style = ParagraphStyle(
        name="DocBullet",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#334155"),
        leftIndent=15,
        spaceAfter=4,
    )
    
    meta_label_style = ParagraphStyle(
        name="MetaLabel",
        fontName="Helvetica-Bold",
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor("#334155"),
    )
    
    meta_val_style = ParagraphStyle(
        name="MetaValue",
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor("#0f172a"),
    )
    
    th_style = ParagraphStyle(
        name="TableHeader",
        fontName="Helvetica-Bold",
        fontSize=9.5,
        leading=12,
        textColor=colors.white,
    )
    
    td_style = ParagraphStyle(
        name="TableCellText",
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#334155"),
    )
    
    story = []
    
    # 1. Title
    story.append(Paragraph(template["title"], title_style))
    story.append(Spacer(1, 10))
    
    # 2. Meeting Information Metadata Table
    story.append(Paragraph("Meeting Information", h1_style))
    story.append(Spacer(1, 4))
    
    metadata_rows = [
        [Paragraph("Date & Time", meta_label_style), Paragraph(template["date"], meta_val_style)],
        [Paragraph("Meeting Type", meta_label_style), Paragraph(template["type"], meta_val_style)],
        [Paragraph("Location", meta_label_style), Paragraph(template["location"], meta_val_style)],
    ]
    if template["participants"]:
        metadata_rows.append([
            Paragraph("Participants", meta_label_style),
            Paragraph(", ".join(template["participants"]), meta_val_style)
        ])
    if template["tags"]:
        metadata_rows.append([
            Paragraph("Tags", meta_label_style),
            Paragraph(", ".join(template["tags"]), meta_val_style)
        ])
        
    meta_table = Table(metadata_rows, colWidths=[120, 375.27])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f8fafc")),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 10))
    
    # 3. Summary Section
    story.append(Paragraph("Summary", h1_style))
    story.append(Spacer(1, 4))
    for line in template["summary"].split("\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), body_style))
    story.append(Spacer(1, 10))
    
    # 4. Action Items Section (Table format)
    if template["actionItems"]:
        story.append(Paragraph("Action Items", h1_style))
        story.append(Spacer(1, 4))
        
        table_data = [[
            Paragraph("Task", th_style),
            Paragraph("Assignee", th_style),
            Paragraph("Due Date", th_style)
        ]]
        
        for item in template["actionItems"]:
            task = clean_markdown(sanitize_text(item.get("task") or ""))
            assignee = clean_markdown(sanitize_text(item.get("assignee") or "Not specified"))
            due_date = clean_markdown(sanitize_text(item.get("dueDate") or "Not specified"))
            
            table_data.append([
                Paragraph(task, td_style),
                Paragraph(assignee, td_style),
                Paragraph(due_date, td_style)
            ])
            
        action_table = Table(table_data, colWidths=[297.16, 99.05, 99.05])
        
        t_style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#059669")), # Emerald 600
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('LEFTPADDING', (0,0), (-1,-1), 8),
            ('RIGHTPADDING', (0,0), (-1,-1), 8),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ])
        for i in range(1, len(table_data)):
            bg_color = colors.HexColor("#f8fafc") if i % 2 == 1 else colors.white
            t_style.add('BACKGROUND', (0, i), (-1, i), bg_color)
            
        action_table.setStyle(t_style)
        story.append(action_table)
        story.append(Spacer(1, 10))
        
    # 5. Key Decisions Section
    if template["decisions"]:
        story.append(Paragraph("Key Decisions", h1_style))
        story.append(Spacer(1, 4))
        for idx, item in enumerate(template["decisions"]):
            decision = clean_markdown(sanitize_text(item.get("decision") or ""))
            context = clean_markdown(sanitize_text(item.get("context") or ""))
            
            decision_text = f"<b>{idx + 1}. {decision}</b>"
            if context:
                decision_text += f"<br/>&nbsp;&nbsp;&nbsp;&nbsp;<i>Context: {context}</i>"
            story.append(Paragraph(decision_text, bullet_style))
        story.append(Spacer(1, 10))
        
    # 6. Deadlines Section (Table format)
    if template["deadlines"]:
        story.append(Paragraph("Deadlines", h1_style))
        story.append(Spacer(1, 4))
        
        table_data = [[
            Paragraph("Item / Task", th_style),
            Paragraph("Deadline Date", th_style),
            Paragraph("Owner", th_style)
        ]]
        
        for item in template["deadlines"]:
            name = clean_markdown(sanitize_text(item.get("item") or ""))
            date_val = clean_markdown(sanitize_text(item.get("date") or "Not specified"))
            owner = clean_markdown(sanitize_text(item.get("owner") or "Not specified"))
            
            table_data.append([
                Paragraph(name, td_style),
                Paragraph(date_val, td_style),
                Paragraph(owner, td_style)
            ])
            
        deadline_table = Table(table_data, colWidths=[297.16, 99.05, 99.05])
        
        t_style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#059669")), # Emerald 600
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('LEFTPADDING', (0,0), (-1,-1), 8),
            ('RIGHTPADDING', (0,0), (-1,-1), 8),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ])
        for i in range(1, len(table_data)):
            bg_color = colors.HexColor("#f8fafc") if i % 2 == 1 else colors.white
            t_style.add('BACKGROUND', (0, i), (-1, i), bg_color)
            
        deadline_table.setStyle(t_style)
        story.append(deadline_table)
        story.append(Spacer(1, 10))
        
    doc.build(story, onFirstPage=draw_first_page, onLaterPages=draw_later_pages)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


def set_cell_background(cell, hex_color: str):
    """Shade a Word table cell with the provided Hex Color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_heading(paragraph, color):
    """Style a Word paragraph as a level-1 heading with custom margins and colors."""
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.color.rgb = color
        run.bold = True


def generate_docx(summary_data: dict[str, Any]) -> bytes:
    """Generate a formatted Word DOCX document matching the PDF layout structure."""
    validated = validate_summary_data(summary_data)
    template = create_meeting_notes_template(validated)
    
    doc = Document()
    
    EMERALD_COLOR = RGBColor(5, 150, 105) # Emerald 600
    SLATE_900_COLOR = RGBColor(15, 23, 42) # Slate 900
    SLATE_700_COLOR = RGBColor(51, 65, 85) # Slate 700
    
    # 1. Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(14)
    run_title = p_title.add_run(template["title"])
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = SLATE_900_COLOR
    run_title.bold = True
    
    # Add running footer
    section = doc.sections[0]
    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_foot = p_footer.add_run("Generated by MeetNote AI")
    run_foot.font.name = 'Arial'
    run_foot.font.size = Pt(9)
    run_foot.font.color.rgb = SLATE_700_COLOR
    
    # 2. Meeting Information Metadata Table
    h = doc.add_heading("Meeting Information", level=1)
    format_heading(h, EMERALD_COLOR)
    
    info_rows = [
        ("Date & Time", template["date"]),
        ("Type", template["type"]),
        ("Location", template["location"]),
    ]
    if template["participants"]:
        info_rows.append(("Participants", ", ".join(template["participants"])))
    if template["tags"]:
        info_rows.append(("Tags", ", ".join(template["tags"])))
        
    meta_table = doc.add_table(rows=len(info_rows), cols=2)
    meta_table.autofit = False
    
    for idx, (label, val) in enumerate(info_rows):
        row = meta_table.rows[idx]
        
        # Label column (1.8 inches)
        cell_0 = row.cells[0]
        cell_0.width = Inches(1.8)
        p0 = cell_0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        run0 = p0.add_run(label)
        run0.bold = True
        run0.font.name = 'Arial'
        run0.font.size = Pt(9.5)
        run0.font.color.rgb = SLATE_700_COLOR
        set_cell_background(cell_0, "F8FAFC")
        
        # Value column (4.2 inches)
        cell_1 = row.cells[1]
        cell_1.width = Inches(4.2)
        p1 = cell_1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        run1 = p1.add_run(val or "Not specified")
        run1.font.name = 'Arial'
        run1.font.size = Pt(9.5)
        run1.font.color.rgb = SLATE_900_COLOR
        set_cell_background(cell_1, "F8FAFC")
        
    # 3. Summary Section
    h = doc.add_heading("Summary", level=1)
    format_heading(h, EMERALD_COLOR)
    
    for line in template["summary"].split("\n"):
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line.strip())
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = SLATE_700_COLOR
            
    # 4. Action Items Section (Table format)
    if template["actionItems"]:
        h = doc.add_heading("Action Items", level=1)
        format_heading(h, EMERALD_COLOR)
        
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        
        # Header cells
        hdr_cells = table.rows[0].cells
        headers_config = [("Task", 3.6), ("Assignee", 1.2), ("Due Date", 1.2)]
        for i, (text, width) in enumerate(headers_config):
            cell = hdr_cells[i]
            cell.width = Inches(width)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell, "059669")
            
        for idx, item in enumerate(template["actionItems"]):
            task = clean_markdown(sanitize_text(item.get("task") or ""))
            assignee = clean_markdown(sanitize_text(item.get("assignee") or "Not specified"))
            due_date = clean_markdown(sanitize_text(item.get("dueDate") or "Not specified"))
            
            row = table.add_row()
            row_cells = row.cells
            
            row_cells[0].width = Inches(3.6)
            p0 = row_cells[0].paragraphs[0]
            p0.paragraph_format.space_before = Pt(4)
            p0.paragraph_format.space_after = Pt(4)
            run0 = p0.add_run(task)
            run0.font.name = 'Arial'
            run0.font.size = Pt(9.5)
            run0.font.color.rgb = SLATE_900_COLOR
            
            row_cells[1].width = Inches(1.2)
            p1 = row_cells[1].paragraphs[0]
            p1.paragraph_format.space_before = Pt(4)
            p1.paragraph_format.space_after = Pt(4)
            run1 = p1.add_run(assignee)
            run1.font.name = 'Arial'
            run1.font.size = Pt(9.5)
            run1.font.color.rgb = SLATE_700_COLOR
            
            row_cells[2].width = Inches(1.2)
            p2 = row_cells[2].paragraphs[0]
            p2.paragraph_format.space_before = Pt(4)
            p2.paragraph_format.space_after = Pt(4)
            run2 = p2.add_run(due_date)
            run2.font.name = 'Arial'
            run2.font.size = Pt(9.5)
            run2.font.color.rgb = SLATE_700_COLOR
            
            bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
            for cell in row_cells:
                set_cell_background(cell, bg_color)
                
    # 5. Key Decisions Section
    if template["decisions"]:
        h = doc.add_heading("Key Decisions", level=1)
        format_heading(h, EMERALD_COLOR)
        
        for idx, item in enumerate(template["decisions"]):
            decision = clean_markdown(sanitize_text(item.get("decision") or ""))
            context = clean_markdown(sanitize_text(item.get("context") or ""))
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(4)
            
            run_idx = p.add_run(f"{idx + 1}. ")
            run_idx.bold = True
            run_idx.font.name = 'Arial'
            run_idx.font.size = Pt(10)
            run_idx.font.color.rgb = SLATE_900_COLOR
            
            run_desc = p.add_run(decision)
            run_desc.font.name = 'Arial'
            run_desc.font.size = Pt(10)
            run_desc.font.color.rgb = SLATE_900_COLOR
            
            if context:
                p_ctx = doc.add_paragraph()
                p_ctx.paragraph_format.left_indent = Inches(0.4)
                p_ctx.paragraph_format.space_after = Pt(6)
                run_ctx = p_ctx.add_run(f"Context: {context}")
                run_ctx.italic = True
                run_ctx.font.name = 'Arial'
                run_ctx.font.size = Pt(9.5)
                run_ctx.font.color.rgb = SLATE_700_COLOR
                
    # 6. Deadlines Section (Table format)
    if template["deadlines"]:
        h = doc.add_heading("Deadlines", level=1)
        format_heading(h, EMERALD_COLOR)
        
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        
        hdr_cells = table.rows[0].cells
        headers_config = [("Item / Task", 3.6), ("Deadline Date", 1.2), ("Owner", 1.2)]
        for i, (text, width) in enumerate(headers_config):
            cell = hdr_cells[i]
            cell.width = Inches(width)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell, "059669")
            
        for idx, item in enumerate(template["deadlines"]):
            name = clean_markdown(sanitize_text(item.get("item") or ""))
            date_val = clean_markdown(sanitize_text(item.get("date") or "Not specified"))
            owner = clean_markdown(sanitize_text(item.get("owner") or "Not specified"))
            
            row = table.add_row()
            row_cells = row.cells
            
            row_cells[0].width = Inches(3.6)
            p0 = row_cells[0].paragraphs[0]
            p0.paragraph_format.space_before = Pt(4)
            p0.paragraph_format.space_after = Pt(4)
            run0 = p0.add_run(name)
            run0.font.name = 'Arial'
            run0.font.size = Pt(9.5)
            run0.font.color.rgb = SLATE_900_COLOR
            
            row_cells[1].width = Inches(1.2)
            p1 = row_cells[1].paragraphs[0]
            p1.paragraph_format.space_before = Pt(4)
            p1.paragraph_format.space_after = Pt(4)
            run1 = p1.add_run(date_val)
            run1.font.name = 'Arial'
            run1.font.size = Pt(9.5)
            run1.font.color.rgb = SLATE_700_COLOR
            
            row_cells[2].width = Inches(1.2)
            p2 = row_cells[2].paragraphs[0]
            p2.paragraph_format.space_before = Pt(4)
            p2.paragraph_format.space_after = Pt(4)
            run2 = p2.add_run(owner)
            run2.font.name = 'Arial'
            run2.font.size = Pt(9.5)
            run2.font.color.rgb = SLATE_700_COLOR
            
            bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
            for cell in row_cells:
                set_cell_background(cell, bg_color)
                
    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()
    stream.close()
    return docx_bytes
